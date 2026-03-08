from __future__ import annotations

import hashlib
import re
from pathlib import Path
from typing import List, Optional

from docx import Document as WordDocument
from langchain_core.documents import Document
from pypdf import PdfReader
from pypdf.errors import PdfReadError

try:
    import fitz  # type: ignore
except ImportError:  # pragma: no cover - optional dependency
    fitz = None

try:
    from PIL import Image
except ImportError:  # pragma: no cover - optional dependency
    Image = None

try:
    import pytesseract
except ImportError:  # pragma: no cover - optional dependency
    pytesseract = None

from app.utils.models import LoadedDocumentSet


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}


def compute_checksum(path: Path) -> str:
    digest = hashlib.sha1()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def build_doc_id(path: Path, checksum: str) -> str:
    return hashlib.sha1(checksum.encode("utf-8")).hexdigest()[:16]


def normalize_text(text: str) -> str:
    lines = [line.strip() for line in text.splitlines()]
    filtered = [line for line in lines if line]
    normalized = "\n".join(filtered)
    normalized = re.sub(r"[ \t]+", " ", normalized)
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    return normalized.strip()


def load_file(
    path: Path,
    checksum: Optional[str] = None,
    collection_name: str = "default",
    enable_ocr: bool = False,
    ocr_language: str = "eng",
) -> LoadedDocumentSet:
    if not path.exists():
        raise FileNotFoundError("File not found: {path}".format(path=path))

    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            "Unsupported file type: {suffix}. Supported: {supported}".format(
                suffix=suffix,
                supported=", ".join(sorted(SUPPORTED_EXTENSIONS)),
            )
        )

    checksum = checksum or compute_checksum(path)
    doc_id = build_doc_id(path, checksum)

    if suffix == ".pdf":
        documents = _load_pdf(path, doc_id, collection_name, enable_ocr=enable_ocr, ocr_language=ocr_language)
    elif suffix == ".docx":
        documents = _load_docx(path, doc_id, collection_name)
    else:
        documents = _load_text(path, doc_id, collection_name)

    if not documents:
        raise ValueError(
            "No text content extracted from {path}. The PDF may be scanned/image-only or use an unsupported encoding.".format(
                path=path
            )
        )

    return LoadedDocumentSet(
        doc_id=doc_id,
        filename=path.name,
        file_path=str(path.resolve()),
        checksum=checksum,
        file_type=suffix.lstrip("."),
        documents=documents,
        collection_name=collection_name,
    )


def _base_metadata(path: Path, doc_id: str, file_type: str, collection_name: str) -> dict:
    return {
        "doc_id": doc_id,
        "filename": path.name,
        "source_path": str(path.resolve()),
        "file_type": file_type,
        "collection_name": collection_name,
    }


def _load_pdf(
    path: Path,
    doc_id: str,
    collection_name: str = "default",
    enable_ocr: bool = False,
    ocr_language: str = "eng",
) -> List[Document]:
    documents = _load_pdf_with_pymupdf(path, doc_id, collection_name)
    if documents:
        return documents

    documents = _load_pdf_with_pypdf(path, doc_id, collection_name)
    if documents:
        return documents

    if enable_ocr:
        return _load_pdf_with_ocr(path, doc_id, collection_name, ocr_language)

    return []


def _load_pdf_with_pypdf(path: Path, doc_id: str, collection_name: str) -> List[Document]:
    base_metadata = _base_metadata(path, doc_id, "pdf", collection_name)
    documents = []

    try:
        reader = PdfReader(str(path))
    except PdfReadError as error:
        raise ValueError("Could not read PDF {path}: {error}".format(path=path, error=error)) from error

    for index, page in enumerate(reader.pages, start=1):
        text = normalize_text(page.extract_text() or "")
        if not text:
            continue
        metadata = dict(base_metadata)
        metadata["page_number"] = index
        metadata["extractor"] = "pypdf"
        documents.append(Document(page_content=text, metadata=metadata))

    return documents


def _load_pdf_with_pymupdf(path: Path, doc_id: str, collection_name: str) -> List[Document]:
    if fitz is None:
        return []

    base_metadata = _base_metadata(path, doc_id, "pdf", collection_name)
    documents = []
    pdf = fitz.open(str(path))

    try:
        for index, page in enumerate(pdf, start=1):
            text = normalize_text(page.get_text("text") or "")
            if not text:
                continue
            metadata = dict(base_metadata)
            metadata["page_number"] = index
            metadata["extractor"] = "pymupdf"
            documents.append(Document(page_content=text, metadata=metadata))
    finally:
        pdf.close()

    return documents


def _load_pdf_with_ocr(path: Path, doc_id: str, collection_name: str, ocr_language: str) -> List[Document]:
    if fitz is None or pytesseract is None or Image is None:
        return []

    base_metadata = _base_metadata(path, doc_id, "pdf", collection_name)
    documents = []
    pdf = fitz.open(str(path))

    try:
        for index, page in enumerate(pdf, start=1):
            pixmap = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0), alpha=False)
            image = Image.frombytes("RGB", [pixmap.width, pixmap.height], pixmap.samples)
            text = normalize_text(pytesseract.image_to_string(image, lang=ocr_language) or "")
            if not text:
                continue
            metadata = dict(base_metadata)
            metadata["page_number"] = index
            metadata["extractor"] = "ocr"
            documents.append(Document(page_content=text, metadata=metadata))
    finally:
        pdf.close()

    return documents


def _load_docx(path: Path, doc_id: str, collection_name: str) -> List[Document]:
    document = WordDocument(str(path))
    base_metadata = _base_metadata(path, doc_id, "docx", collection_name)
    parts = []
    current_heading = None

    for paragraph in document.paragraphs:
        text = normalize_text(paragraph.text)
        if not text:
            continue

        style_name = getattr(paragraph.style, "name", "") or ""
        if style_name.lower().startswith("heading"):
            current_heading = text
            parts.append("\n{heading}\n".format(heading=text))
            continue

        parts.append(text)

    content = normalize_text("\n".join(parts))
    if not content:
        return []

    metadata = dict(base_metadata)
    metadata["section_heading"] = current_heading
    return [Document(page_content=content, metadata=metadata)]


def _load_text(path: Path, doc_id: str, collection_name: str) -> List[Document]:
    base_metadata = _base_metadata(path, doc_id, path.suffix.lower().lstrip("."), collection_name)
    content = normalize_text(path.read_text(encoding="utf-8"))
    if not content:
        return []
    return [Document(page_content=content, metadata=base_metadata)]
